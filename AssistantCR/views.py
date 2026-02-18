####################    POUR RECUPERER LE LOGO - SUPERFLU ?     ##########################


from django.conf import settings
from django.contrib.staticfiles import finders
from django.http import HttpResponse
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import os
import json

# Constante pour la largeur du logo
LOGO_WIDTH = Inches(2.5)

def find_logo_path():
    """Recherche le chemin du logo dans les répertoires statiques."""
    logo_path = finders.find('logo-massy.png')
    if not logo_path:
        if getattr(settings, 'STATIC_ROOT', None):
            candidate = os.path.join(settings.STATIC_ROOT, 'logo-massy.jpg')
            if os.path.exists(candidate):
                logo_path = candidate
        if not logo_path and getattr(settings, 'STATICFILES_DIRS', None):
            for directory in settings.STATICFILES_DIRS:
                candidate = os.path.join(directory, 'logo-massy.png')
                if os.path.exists(candidate):
                    logo_path = candidate
                    break
    return logo_path


####################    UPLOAD D'UN JSON de CR     ##########################

#from django.shortcuts import render
#from django.contrib import messages
#import json

#def upload_json_file(request):
#    if request.method == 'POST' and request.FILES.get('json_file'):
#        try:
#json_file = request.FILES['json_file']
#            data = json.load(json_file)
#            messages.success(request, 'File uploaded successfully!')
#            return render(request, 'edit_meeting_minutes.html', {'data': data})
#        except json.JSONDecodeError:
#            messages.error(request, 'Invalid JSON file')
#    return render(request, 'upload.html')








####################    MODIFICATION D'UN JSON de CR ET GENERATION DU DOC WORD     ##########################


import json
import os
import re
from datetime import datetime

from django.http import HttpResponse
from django.shortcuts import render
from django.views.decorators.http import require_http_methods

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT  # WD_PARAGRAPH_ALIGNMENT est conservé car déjà utilisé dans ton code
from docx.shared import Pt


# Données factice lorsque l'utilisateur affiche le formulaire directement
def make_demo_meeting():
    return {
        "reunion": {
            "objet": "Réunion fictive pour l'exemple...",
            "date": "2025-09-01",  # format YYYY-MM-DD
            "lieu": "Salle Faro",
            "ordre_du_jour": [
                "Point budgétaire",
                "Travaux en cours",
                "Questions diverses"
            ],
            "participants": [
                {"nom": "Jean Dupont", "role": "Directeur de projet"},
                {"nom": "Marie Martin", "role": "Cheffe de service"},
            ],
            "teneur_de_la_reunion": [
                {
                    "point_ordre_du_jour": "Point budgétaire",
                    "discussion": "Présentation des dépenses et arbitrages. Tout va bien même si ça irait mieux avec plus d'argent."
                },
                {
                    "point_ordre_du_jour": "Travaux en cours",
                    "discussion": "Peinture et réfection des sols, c'est fait. Le reste, ça attendra."
                }
            ],
            "recapitulatif_qui_fait_quoi": [
                {
                    "tache": "Contacter le prestataire",
                    "responsable": "Marie Martin",
                    "echeance": "2025-09-14"
                }
            ]
        }
    }


@require_http_methods(["GET", "POST"])
def generate_meeting_minutes_from_form(request):
    """GET  -> affiche un formulaire vierge (template)
       POST -> génère le .docx à partir d'un JSON (json_data) OU d'un POST classique (champs)."""

    def format_french_date(date_str):
        """YYYY-MM-DD -> '7 mai 2025' (tolérant si vide ou invalide)."""
        if not date_str:
            return "Date non spécifiée"
        try:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            month_names = [
                "janvier", "février", "mars", "avril", "mai", "juin",
                "juillet", "août", "septembre", "octobre", "novembre", "décembre",
            ]
            return f"{dt.day} {month_names[dt.month - 1]} {dt.year}"
        except (ValueError, TypeError):
            return str(date_str)

    def _split_lines(val):
        return [ln.strip() for ln in (val or "").splitlines() if ln.strip()]

    def _parse_participants(val):
        """Chaque ligne: 'Nom (rôle)' ou 'Nom - rôle' ou juste 'Nom'."""
        participants = []
        for line in _split_lines(val):
            m = re.match(r"^\s*(.+?)\s*(?:\((.+?)\)|-(.+?))?\s*$", line)
            if m:
                nom = (m.group(1) or "").strip()
                role = (m.group(2) or m.group(3) or "").strip()
            else:
                nom, role = line, ""
            participants.append({"nom": nom, "role": role})
        return participants

    def _parse_tasks(post):
        """Cherche des groupes action_N / responsable_N / echeance_N dans le POST."""
        tasks = []
        # collecte des index présents
        idxs = set()
        for k in post.keys():
            m = re.match(r"^(?:action|responsable|echeance)_(\d+)$", k)
            if m:
                idxs.add(int(m.group(1)))
        for i in sorted(idxs):
            t = (post.get(f"action_{i}", "") or "").strip()
            r = (post.get(f"responsable_{i}", "") or "").strip()
            e = (post.get(f"echeance_{i}", "") or "").strip()
            if t or r or e:
                tasks.append({"tache": t, "responsable": r, "echeance": e})
        return tasks

    def _build_reunion_from_form(post):
        """Construit la structure 'reunion' depuis un POST non-JSON."""
        objet = (post.get("objet") or post.get("titre") or "Réunion").strip()
        date  = (post.get("date") or "").strip()
        lieu  = (post.get("lieu") or "").strip()

        ordre_du_jour = _split_lines(post.get("ordre_du_jour"))
        participants  = _parse_participants(post.get("participants"))

        # Teneur: soit liste de points (point_1, discussion_1), soit un bloc 'teneur'
        points = []
        idxs = set()
        for k in post.keys():
            m = re.match(r"^(?:point|discussion)_(\d+)$", k)
            if m:
                idxs.add(int(m.group(1)))
        for i in sorted(idxs):
            p_title = (post.get(f"point_{i}", "") or "").strip()
            p_disc  = (post.get(f"discussion_{i}", "") or "").strip()
            if p_title or p_disc:
                points.append({
                    "point_ordre_du_jour": p_title,
                    "discussion": p_disc,
                })

        # fallback: un seul bloc 'teneur' libre
        if not points and (post.get("teneur") or "").strip():
            points = [{
                "point_ordre_du_jour": "Points discutés",
                "discussion": (post.get("teneur") or "").strip(),
            }]

        recap = _parse_tasks(post)

        return {
            "objet": objet,
            "date": date,
            "lieu": lieu or "Lieu non spécifié",
            "ordre_du_jour": ordre_du_jour,
            "participants": participants,
            "teneur_de_la_reunion": points,
            "recapitulatif_qui_fait_quoi": recap,
        }

    # --- GET: affiche le formulaire vierge ---
    if request.method == "GET":
        skeleton = make_demo_meeting()
        return render(
            request,
            "AssistantCR/edit_meeting_minutes.html",
            {
                "data": skeleton,                                  # si ton template lit un dict
                "json_data": json.dumps(skeleton, ensure_ascii=False),  # si ton bouton export attend une string JSON
            },
        )

    # --- POST: JSON ou POST classique ---
    try:
        reunion = None
        json_data = request.POST.get("json_data")
        if json_data:
            data = json.loads(json_data)
            reunion = data.get("reunion")
            if not reunion:
                return HttpResponse("Format JSON invalide.", status=400)
        else:
            reunion = _build_reunion_from_form(request.POST)

        # --- Génération du document ---
        document = Document()

        # Style dédié pour la teneur (justifié)
        try:
            teneur_style = document.styles.add_style('TeneurTexte', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            # Style déjà existant
            teneur_style = document.styles['TeneurTexte']
        teneur_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Logo (optionnel)
        try:
            logo_path = find_logo_path()
        except Exception:
            logo_path = None
        if logo_path and os.path.exists(logo_path):
            p = document.add_paragraph()
            run = p.add_run()
            run.add_picture(logo_path, width=LOGO_WIDTH)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_paragraph()  # espace après logo

        # Titre
        heading = document.add_heading(level=1)
        heading.add_run(reunion.get("objet", "Réunion"))
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Sous-titre date/lieu
        formatted_date = format_french_date(reunion.get("date"))
        location = reunion.get("lieu", "Lieu non spécifié")
        sub = document.add_heading(level=2)
        sub.add_run(f"Réunion du {formatted_date} tenue à {location}")
        sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph()

        # Ordre du jour
        document.add_heading("Ordre du Jour", level=2)
        if reunion.get("ordre_du_jour"):
            for item in reunion["ordre_du_jour"]:
                document.add_paragraph(item, style="List Bullet")
        else:
            document.add_paragraph("Aucun point à l'ordre du jour.")

        # Participants
        document.add_heading("Participants", level=2)
        if reunion.get("participants"):
            for participant in reunion["participants"]:
                nom = participant.get("nom", "Non spécifié")
                role = participant.get("role", "")
                label = f"{nom} ({role})" if role else nom
                document.add_paragraph(label, style="List Bullet")
        else:
            document.add_paragraph("Aucun participant enregistré.")

        # Teneur (JUSTIFIÉE uniquement pour les points discutés)
        document.add_heading("Teneur de la Réunion", level=2)
        if reunion.get("teneur_de_la_reunion"):
            for point in reunion["teneur_de_la_reunion"]:
                if point.get("point_ordre_du_jour"):
                    document.add_heading(point["point_ordre_du_jour"], level=3)
                if point.get("discussion"):
                    p = document.add_paragraph(point["discussion"])
                    # Applique le style justifié
                    p.style = 'TeneurTexte'
                    # Et force l'alignement au cas où un style global écraserait
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p = document.add_paragraph("Aucun point discuté enregistré.")
            p.style = 'TeneurTexte'
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Récapitulatif
        document.add_paragraph()
        document.add_heading("Récapitulatif Qui Fait Quoi", level=2)
        document.add_paragraph()

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Action"
        header_cells[1].text = "Responsable"
        header_cells[2].text = "Date cible"

        # Gras en-têtes
        for cell in header_cells:
            for r in cell.paragraphs[0].runs:
                r.bold = True

        tasks = reunion.get("recapitulatif_qui_fait_quoi") or []
        if tasks:
            for task in tasks:
                row = table.add_row().cells
                row[0].text = task.get("tache", "")
                row[1].text = task.get("responsable", "")
                row[2].text = format_french_date(task.get("echeance", ""))
        else:
            row = table.add_row().cells
            row[0].text = "Aucune action enregistrée"
            row[1].text = ""
            row[2].text = ""

        # Note finale
        document.add_paragraph()
        sep = document.add_paragraph("_____________________________________________________________")
        sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph()
        note = document.add_paragraph()
        r = note.add_run("Ce document a été généré avec l'aide de l'IA Mistral - Massy Innove version 1.1")
        r.italic = True
        r.font.size = Pt(8)
        note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Réponse .docx
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = "attachment; filename=compte_rendu_reunion.docx"
        document.save(response)
        return response

    except json.JSONDecodeError:
        return HttpResponse("Erreur de décodage JSON.", status=400)
    except Exception as e:
        return HttpResponse(f"Une erreur est survenue: {e}", status=500)


####################### GENERATION DU PDF

# generate_meeting_minutes_pdf.py

# views.py
import os, json
from datetime import datetime
from django.http import HttpResponse
from django.shortcuts import render
from django.views.decorators.http import require_http_methods

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
)



@require_http_methods(["GET", "POST"])
def generate_meeting_minutes_pdf(request):
    """Génère un PDF fidèle au style Word, sans dépendance externe."""

    # --- Format date ---
    def format_french_date(date_str):
        if not date_str:
            return "Date non spécifiée"
        try:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            mois = [
                "janvier","février","mars","avril","mai","juin",
                "juillet","août","septembre","octobre","novembre","décembre"
            ]
            return f"{dt.day} {mois[dt.month-1]} {dt.year}"
        except Exception:
            return str(date_str)

    # --- Polices système disponibles (urw-base35) ---
    # --- Polices PDF de base (stables, intégrées) ---
    body_font = "Times-Roman"       # équivalent Cambria (Corps)
    body_italic = "Times-Italic"
    title_font = "Helvetica-Bold"   # équivalent Calibri Bold
    subtitle_font = "Helvetica-Bold"


    # --- Styles fidèles au Word ---
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="TitrePrincipal",
        fontName=title_font,
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#365F92"),
        alignment=TA_CENTER,
        spaceBefore=12,
        spaceAfter=10,
    ))

    styles.add(ParagraphStyle(
        name="SousTitre",
        fontName=subtitle_font,
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#4F82BD"),
        alignment=TA_CENTER,
        spaceBefore=4,
        spaceAfter=14,
    ))

    styles.add(ParagraphStyle(
        name="TitreSection",
        fontName=title_font,
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#4F82BD"),
        alignment=TA_LEFT,
        spaceBefore=16,
        spaceAfter=8,
    ))

    styles.add(ParagraphStyle(
        name="TitreSousSection",
        fontName=title_font,
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#4F82BD"),
        alignment=TA_LEFT,
        spaceBefore=10,
        spaceAfter=4,
    ))

    styles.add(ParagraphStyle(
        name="TexteNormal",
        fontName=body_font,
        fontSize=11,
        leading=14.5,
        textColor=colors.black,
        alignment=TA_JUSTIFY,
        spaceBefore=0,
        spaceAfter=8,
    ))

    styles.add(ParagraphStyle(
        name="NoteFinale",
        fontName=body_italic,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#555555"),
        alignment=TA_CENTER,
        spaceBefore=18,
    ))

    # --- GET : affiche formulaire ---
    if request.method == "GET":
        skeleton = make_demo_meeting()
        return render(request, "AssistantCR/edit_meeting_minutes.html", {
            "data": skeleton,
            "json_data": json.dumps(skeleton, ensure_ascii=False),
        })

    # --- POST : génération du PDF ---
    try:
        data = json.loads(request.POST.get("json_data", "{}"))
        reunion = data.get("reunion", {})

        story = []

        # --- Logo (optionnel)
        try:
            logo_path = find_logo_path()
        except Exception:
            logo_path = None
        if logo_path and os.path.exists(logo_path):
            img = Image(logo_path)
            iw, ih = img.imageWidth, img.imageHeight
            if iw and ih:
                img.drawWidth = 5.5 * cm
                img.drawHeight = img.drawWidth * ih / iw
            img.hAlign = "CENTER"
            story += [img, Spacer(1, 0.6 * cm)]

        # --- Titre principal + sous-titre
        story.append(Paragraph(reunion.get("objet", "Réunion"), styles["TitrePrincipal"]))
        story.append(Paragraph(
            f"Réunion du {format_french_date(reunion.get('date'))} tenue à {reunion.get('lieu','Lieu non spécifié')}",
            styles["SousTitre"]
        ))
        story.append(Spacer(1, 0.3 * cm))

        # --- Ordre du jour
        story.append(Paragraph("Ordre du Jour", styles["TitreSection"]))
        for item in reunion.get("ordre_du_jour", []):
            story.append(Paragraph(f"• {item}", styles["TexteNormal"]))
        if not reunion.get("ordre_du_jour"):
            story.append(Paragraph("Aucun point à l'ordre du jour.", styles["TexteNormal"]))

        # --- Participants
        story.append(Paragraph("Participants", styles["TitreSection"]))
        participants = reunion.get("participants", [])
        if participants:
            for p in participants:
                nom, role = p.get("nom",""), p.get("role","")
                story.append(Paragraph(f"• {nom} ({role})" if role else f"• {nom}", styles["TexteNormal"]))
        else:
            story.append(Paragraph("Aucun participant enregistré.", styles["TexteNormal"]))

        # --- Teneur
        story.append(Paragraph("Teneur de la Réunion", styles["TitreSection"]))
        for point in reunion.get("teneur_de_la_reunion", []):
            if point.get("point_ordre_du_jour"):
                story.append(Paragraph(point["point_ordre_du_jour"], styles["TitreSousSection"]))
            if point.get("discussion"):
                for para in point["discussion"].splitlines():
                    story.append(Paragraph(para.strip(), styles["TexteNormal"]))
            story.append(Spacer(1, 0.15 * cm))

        # --- Récapitulatif
        story.append(Paragraph("Récapitulatif Qui Fait Quoi", styles["TitreSection"]))
        data_rows = [["Action", "Responsable", "Date cible"]]
        for t in reunion.get("recapitulatif_qui_fait_quoi", []):
            data_rows.append([
                t.get("tache",""),
                t.get("responsable",""),
                format_french_date(t.get("echeance","")) if t.get("echeance") else ""
            ])
        if len(data_rows) == 1:
            data_rows.append(["Aucune action enregistrée", "", ""])

        tbl = Table(data_rows, colWidths=[9*cm, 5*cm, 3*cm])
        tbl.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#f3f6f6")),
            ("FONTNAME", (0,0), (-1,0), title_font),
            ("FONTNAME", (0,1), (-1,-1), body_font),
            ("FONTSIZE", (0,0), (-1,-1), 10.5),
            ("ALIGN", (0,0), (-1,0), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(tbl)

        # --- Ligne + note finale
        story += [
            Spacer(1, 0.5 * cm),
            Paragraph("_____________________________________________________________", styles["SousTitre"]),
            Paragraph("Ce document a été généré avec l'aide de l'IA Mistral - Massy Innove version 1.1", styles["NoteFinale"]),
        ]

        # --- Génération PDF
        response = HttpResponse(content_type="application/pdf")
        response["Content-Disposition"] = 'attachment; filename="compte_rendu_reunion.pdf"'
        doc = SimpleDocTemplate(
            response,
            pagesize=A4,
            leftMargin=2.54 * cm,
            rightMargin=2.54 * cm,
            topMargin=2.54 * cm,
            bottomMargin=2.54 * cm,
        )
        doc.build(story)
        return response

    except Exception as e:
        return HttpResponse(f"Erreur PDF : {e}", status=500)



####################    CREATION D'UN JSON de CR à partir d'une transcription traitée par IA     ##########################


from django.conf import settings
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.views.decorators.http import require_http_methods
import io
import json
from mistralai import Mistral

# Constants for the Mistral completion configuration
MODEL = "mistral-large-latest"
TEMPERATURE = 0.7


def generate_response_format_and_instructions():
    # Définition centralisée de la structure
    json_schema = {
        "type": "object",
        "properties": {
            "reunion": {
                "type": "object",
                "properties": {
                    "objet": {"type": "string"},
                    "date": {
                        "type": "string",
                        "description": "Format YYYY-MM-DD, par exemple 2025-05-07"
                    },
                    "lieu": {"type": "string"},
                    "ordre_du_jour": {
                        "type": "array",
                        "items": {"type": "string"}
                    },
                    "participants": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "nom": {"type": "string"},
                                "role": {"type": "string"}
                            },
                            "required": ["nom", "role"]
                        }
                    },
                    "teneur_de_la_reunion": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "point_ordre_du_jour": {"type": "string"},
                                "discussion": {"type": "string"}
                            },
                            "required": ["point_ordre_du_jour", "discussion"]
                        }
                    },
                    "recapitulatif_qui_fait_quoi": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "tache": {"type": "string"},
                                "responsable": {"type": "string"},
                                "echeance": {
                                    "type": "string",
                                    "description": "Format YYYY-MM-DD, par exemple 2025-05-07"
                                }
                            },
                            "required": ["tache", "responsable", "echeance"]
                        }
                    }
                },
                "required": [
                    "objet",
                    "date",
                    "lieu",
                    "ordre_du_jour",
                    "participants",
                    "teneur_de_la_reunion",
                    "recapitulatif_qui_fait_quoi"
                ]
            }
        },
        "required": ["reunion"]
    }

    # Génération des instructions textuelles
    instructions = f"""
Le texte envoyé par l'utilisateur est la transcription d'une réunion. Structure ces informations selon le format JSON suivant :

{json.dumps(json_schema, indent=2)}

RÈGLES STRICTES À SUIVRE ABSOLUMENT :
1. La réponse DOIT être un JSON valide et rien d'autre
2. Toutes les dates (date principale et échéances) DOIVENT être au format YYYY-MM-DD sans exception
   - Exemple : "2025-05-07" pour le 7 mai 2025
   - Si la date est mentionnée sous une autre forme (comme "7 mai 2025"), tu DOIS la convertir en YYYY-MM-DD
   - Si aucune date n'est mentionnée, utilise la date du jour au format YYYY-MM-DD
3. Pour les échéances dans le récapitulatif :
   - Si une date est mentionnée dans le texte, utilise-la en la convertissant au format YYYY-MM-DD
   - Si aucune date n'est mentionnée, utilise une date raisonnable (par exemple 1 mois après la date de réunion)
   - Ne laisse jamais une date vide - utilise au moins la date du jour si aucune information
4. Pour les participants :
   - Si le rôle est mentionné entre parenthèses comme "Nom (Rôle)", extraire "Nom" et "Rôle"
   - Si seulement un nom est donné, laisser le rôle vide
5. Tous les champs requis doivent être présents même si vides
6. Les points de l'ordre du jour doivent correspondre exactement à ceux dans la teneur de la réunion
7. Lorsque tu mentionnes des personnes avec leur nom et prénom, le format à utiliser est "Prénom NOM" et non pas "NOM Prénom" même si la transcription t'indique les noms ainsi.
8. Les points à l'ordre du jour doivent être résumés ("point_ordre_du_jour" dans le JSON) mais la discussion attachée à chacun de ces points ("discussion" dans le JSON) doit être suffisamment détaillée pour reflêter l'intégralité des échanges.
"""

    return {
        'response_format': {
            "type": "json_object",
            "json_schema": json_schema
        },
        'instructions': instructions
    }

# Configuration centrale
config = generate_response_format_and_instructions()
RESPONSE_FORMAT = config['response_format']
BASIC_INSTRUCTIONS = config['instructions']


####################### La vue qui enchaîne tout




from django.views.decorators.http import require_http_methods
from django.http import HttpResponse
from django.shortcuts import render
import io, json
from docx import Document
from mistralai import Mistral
from django.conf import settings
import time
from core.ai_audit.services import log_ai_call, hash_request
from django.db import connection


@require_http_methods(["GET", "POST"])
def upload_and_process(request):
    if request.method == 'POST':
        # --- Variables de logging (n’influencent pas le flux métier)
        status_for_log = "success"
        input_tokens = 0
        output_tokens = 0
        latency_ms = 0
        model_name = MODEL  # même constante que pour l'appel
        method = request.POST.get('upload_method')
        text_for_hash = ""  # on ne hashera qu'en finally

        try:
            # Déterminer la méthode d'upload utilisée
            text_content = ""

            # Gestion des différentes méthodes d'upload
            if method == 'paste':
                # Méthode 1: Texte collé directement
                text_content = request.POST.get('pasted_text', '')
                if not text_content.strip():
                    return HttpResponse("Erreur : Le texte collé est vide.", status=400)

            elif method == 'txt':
                # Méthode 2: Fichier TXT (méthode actuelle)
                txt_file = request.FILES.get('txt_file')
                if not txt_file:
                    return HttpResponse("Erreur : Aucun fichier TXT n'a été envoyé.", status=400)

                # Lecture du fichier avec différents encodages
                encodings = ['utf-8', 'ISO-8859-1', 'Windows-1252']
                for encoding in encodings:
                    try:
                        text_content = txt_file.read().decode(encoding)
                        break
                    except UnicodeDecodeError:
                        txt_file.seek(0)  # Réinitialiser le pointeur de fichier

                if text_content is None:
                    return HttpResponse("Erreur : Impossible de décoder le fichier TXT.", status=400)
                if not text_content.strip():
                    return HttpResponse("Erreur : Le contenu du fichier TXT est vide.", status=400)

            elif method == 'word':
                # Méthode 3: Fichier Word
                word_file = request.FILES.get('word_file')
                if not word_file:
                    return HttpResponse("Erreur : Aucun fichier Word n'a été envoyé.", status=400)

                try:
                    doc = Document(io.BytesIO(word_file.read()))
                    text_content = "\n".join([para.text for para in doc.paragraphs])
                    if not text_content.strip():
                        return HttpResponse("Erreur : Le contenu du fichier Word est vide.", status=400)
                except Exception as e:
                    return HttpResponse(f"Erreur lors de la lecture du fichier Word: {str(e)}", status=400)

            else:
                return HttpResponse("Erreur : Méthode d'upload non valide.", status=400)

            # Copie pour le hash (on NE modifie PAS text_content)
            text_for_hash = (text_content or "")[:]

            # Appel à l'API Mistral — inchangé
            client = Mistral(api_key=settings.MISTRAL_API_KEY)
            messages = [
                {"role": "system", "content": BASIC_INSTRUCTIONS},
                {"role": "user", "content": text_content},
            ]

            t0 = time.perf_counter()
            try:
                # Appel à l'API avec gestion des erreurs
                chat_response = client.chat.complete(
                    model=MODEL,
                    messages=messages,
                    temperature=TEMPERATURE,
                    response_format=RESPONSE_FORMAT
                )
            finally:
                latency_ms = int((time.perf_counter() - t0) * 1000)

            # Extraction des tokens si fournis par le SDK
            try:
                u = getattr(chat_response, "usage", None)
                if u is not None:
                    input_tokens = int(getattr(u, "prompt_tokens", 0) or 0)
                    output_tokens = int(getattr(u, "completion_tokens", 0) or 0)
                elif isinstance(chat_response, dict):
                    uu = chat_response.get("usage", {}) or {}
                    input_tokens = int(uu.get("prompt_tokens", 0) or 0)
                    output_tokens = int(uu.get("completion_tokens", 0) or 0)
            except Exception:
                input_tokens = output_tokens = 0  # ne pas casser le flux

            # Vérification de la réponse — inchangé
            if not chat_response or not chat_response.choices:
                status_for_log = "error"
                return HttpResponse("Erreur : Pas de réponse de l'API Mistral.", status=500)
            if not hasattr(chat_response.choices[0].message, 'content'):
                status_for_log = "error"
                return HttpResponse("Erreur : La réponse de Mistral n'a pas de contenu.", status=500)

            # Traitement de la réponse — inchangé
            json_content = chat_response.choices[0].message.content
            debug_info = f"<h2>Débogage</h2><p>JSON brut reçu :</p><pre>{json_content}</pre>"

            try:
                data = json.loads(json_content)
                debug_info += f"<p>Données parsées :</p><pre>{json.dumps(data, indent=2)}</pre>"

                if isinstance(data, dict):
                    debug_info += "<p>Clés disponibles :</p><ul>"
                    for key in data.keys():
                        debug_info += f"<li>{key}</li>"
                    debug_info += "</ul>"

                    if 'reunion' in data:
                        debug_info += "<p>Structure reunion trouvée</p>"
                        reunion_data = data['reunion']
                        if isinstance(reunion_data, dict):
                            debug_info += "<p>Clés dans reunion :</p><ul>"
                            for key in reunion_data.keys():
                                debug_info += f"<li>{key}</li>"
                            debug_info += "</ul>"
                    else:
                        status_for_log = "error"
                        return HttpResponse(f"Erreur : La structure JSON ne contient pas 'reunion'. {debug_info}", status=400)

                return render(
                    request,
                    "AssistantCR/edit_meeting_minutes.html",
                    {
                        "data": data,                                   # dict parsé depuis Mistral
                        "json_data": json.dumps(data, ensure_ascii=False),
                        "debug_info": debug_info,                       # optionnel
                    },
                )

            except json.JSONDecodeError as e:
                status_for_log = "error"
                return HttpResponse(f"Erreur de décodage JSON: {str(e)}<br><br>Contenu problématique :<pre>{json_content}</pre>", status=400)

        except Exception as e:
            status_for_log = "error"
            return HttpResponse(f"Erreur inattendue: {str(e)}", status=500)

        finally:
            # Logging découplé (jamais bloquant)
            try:
                connection.close()
                log_ai_call(
                    user=request.user if request.user.is_authenticated else None,
                    source_app="AssistantCR",                 # nom de cette app (à ajuster si besoin)
                    source_module="views.upload_and_process",
                    provider="mistral",
                    model=model_name,                         # ex. "mistral-large-latest"
                    tokens_input=input_tokens,
                    tokens_output=output_tokens,
                    request_hash=hash_request(text_for_hash) if text_for_hash else "",
                    latency_ms=latency_ms,
                    status=status_for_log,
                    metadata={
                        "upload_method": method,
                        "response_has_usage": bool(input_tokens or output_tokens),
                    },
                    # cost_eur non fourni -> calcul auto via MistralModel (par 1M) + 1,5 mg/token
                )
            except Exception:
                pass

    # Pour les requêtes GET, afficher simplement le formulaire
    return render(request, 'upload.html')



def upload_form(request):
    return render(request, 'AssistantCR/home.html')
    
    
    